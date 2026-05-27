#!/usr/bin/env python3
"""Build an editable .docx version of Mathias's music-store résumé."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path("/Users/victoria/Desktop/Mathias-Resume-MusicStore.docx")

doc = Document()

# tight margins
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

# default body font
style = doc.styles["Normal"]
style.font.name = "Garamond"
style.font.size = Pt(11)

ACCENT = RGBColor(0x8B, 0x00, 0x00)
MUTED = RGBColor(0x6B, 0x72, 0x80)

def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Helvetica Neue"
    r.font.size = Pt(22)

def contact(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.name = "Helvetica Neue"
    r.font.size = Pt(10.5)
    r.font.color.rgb = ACCENT

def body(text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

# ── Header ─────────────────────────────────────────
h1("MATHIAS LITTLEJOHN")
contact("Wellington, Florida  ·  (561) 906-8731  ·  Mathiasrigolittlejohn@gmail.com")

# ── Objective ──────────────────────────────────────
h2("Objective")
body("A part-time or summer position with a Wellington-area music store where I can be useful "
     "around the floor, the gear, and the players — and keep learning from people who know more "
     "than I do.")

# ── Music ──────────────────────────────────────────
h2("Music")
bullet(" — primary instrument; comfortable with electric and acoustic.", bold_lead="Guitar")
bullet(" — second instrument; concert-band reading experience.", bold_lead="Saxophone")
bullet(" — for fun, mostly blues and classic-rock grooves.", bold_lead="Harmonica")
bullet(" deep love of classic rock and everything around it — blues, Southern rock, "
       "'70s singer-songwriter, early metal, alt-rock that grew out of it.", bold_lead="Genres:")
bullet(" regular at Wellington-area open mics, including Village Music Wellington.",
       bold_lead="Live experience:")

# ── Why I'd Be Useful ──────────────────────────────
h2("Why I'd Be Useful on the Floor")
bullet("I can actually demo a guitar for a customer.")
bullet("Comfortable explaining gear in plain English.")
bullet("Patient with first-time buyers and parents.")
bullet("Careful with inventory and consignment instruments.")
bullet("Reliable, on time, dressed for the room.")
bullet("Quick to learn POS systems, repair intake, restringing.")

# ── Languages ──────────────────────────────────────
h2("Languages")
body("Fluent in English, Portuguese, and Spanish.")

# ── Experience ─────────────────────────────────────
h2("Experience")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
r = p.add_run("Volunteer Assistant")
r.bold = True
p.add_run("  ·  Bill Baggs Cape Florida State Park — Key Biscayne, FL · Summer 2025. "
          "Public-facing role assisting park staff and visitors.")

# ── Athletics ──────────────────────────────────────
h2("Athletics")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
r = p.add_run("Competitive Club Swimmer")
r.bold = True
p.add_run(", Wellington FL (2024–Present)  ·  ")
r = p.add_run("Varsity Swim Team")
r.bold = True
p.add_run(", Wellington High School. Daily training = discipline and time management.")

# ── Education ──────────────────────────────────────
h2("Education")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Wellington High School")
r.bold = True
p.add_run("  ·  Class of 2029  ·  Cambridge AICE Diploma track  ·  Honor roll.")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
r = p.add_run("Aubrick School")
r.bold = True
p.add_run(", São Paulo, Brazil (2016–2023)  ·  Student of the Week, Felsted International Summer School 2022.")

# ── Roles of Interest ──────────────────────────────
h2("Roles of Interest")
body("Sales floor associate · gear demo / setup helper · lesson-program front desk · "
     "restringing & basic setup tech-in-training · open-mic night support · "
     "inventory & receiving · anything that gets me near the gear and the customers.")

# ── References ─────────────────────────────────────
h2("References")
body("Available upon request.")

doc.save(OUT)
print(f"WROTE {OUT}")
